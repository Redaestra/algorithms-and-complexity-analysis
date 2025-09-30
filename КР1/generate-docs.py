from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(14)

def add_title_page():
    doc.add_paragraph('Министерство науки и высшего образования Российской Федерации').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph('[Название вуза]').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph('[Факультет / Институт]').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph('[Кафедра]').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph('\n' * 4)

    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run('ОТЧЁТ\nпо выполнению контрольной работы №1\nпо дисциплине «Алгоритмы и анализ сложности»')
    run.bold = True

    doc.add_paragraph('\n' * 4)
    doc.add_paragraph('Студент: Фамилия Имя')
    doc.add_paragraph('Группа: [Номер группы]')
    doc.add_paragraph('Преподаватель: [ФИО преподавателя]')
    doc.add_paragraph('Дата: [Дата сдачи]')
    doc.add_paragraph('\n' * 4)
    doc.add_paragraph('Город, 2025').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_page_break()

def add_table_of_contents(tasks):
    doc.add_heading('Содержание', level=1)
    for task_num in tasks:
        p = doc.add_paragraph()
        p.add_run(f'Задача {task_num}').bold = True
    doc.add_page_break()

def add_task_section(task_num):
    heading = doc.add_heading(level=2)
    run = heading.add_run(f'Задача {task_num}')
    bookmark = f'задача-{task_num}'
    tag = heading._element
    start = OxmlElement('w:bookmarkStart')
    start.set(qn('w:id'), '0')
    start.set(qn('w:name'), bookmark)
    tag.insert(0, start)
    end = OxmlElement('w:bookmarkEnd')
    end.set(qn('w:id'), '0')
    tag.append(end)
    doc.add_paragraph('Условие:\n(здесь указывается условие задачи)')
    doc.add_paragraph('Решение:\n(решение задачи с пояснениями)')
    doc.add_page_break()

tasks = [3, 5, 8, 11, 14, 16, 20, 23, 25, 30]
add_title_page()
add_table_of_contents(tasks)
for task in tasks:
    add_task_section(task)

doc.save("Контрольная_работа_1_Алгоритмы.docx")
